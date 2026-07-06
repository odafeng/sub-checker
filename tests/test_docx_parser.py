"""Tests for .docx parser."""

from pathlib import Path

from sub_checker.parsers.docx_parser import parse_docx


def test_parse_docx_sections(sample_docx: Path, sample_figures_dir: Path):
    ms = parse_docx(sample_docx, sample_figures_dir)
    headings = [s.heading for s in ms.sections]
    assert "Abstract" in headings
    assert "Introduction" in headings
    assert "Methods" in headings
    assert "Results" in headings
    assert "Discussion" in headings
    assert "References" in headings


def test_parse_docx_paragraphs(sample_docx: Path, sample_figures_dir: Path):
    ms = parse_docx(sample_docx, sample_figures_dir)
    assert len(ms.paragraphs) > 0
    assert len(ms.raw_text) > 0


def test_parse_docx_references(sample_docx: Path, sample_figures_dir: Path):
    ms = parse_docx(sample_docx, sample_figures_dir)
    assert ms.reference_section is not None
    assert "Smith" in ms.reference_section
    assert "Jones" in ms.reference_section


def test_parse_docx_figure_dir(sample_docx: Path, sample_figures_dir: Path):
    ms = parse_docx(sample_docx, sample_figures_dir)
    assert ms.figure_dir == sample_figures_dir


def test_parse_docx_title(sample_docx: Path, sample_figures_dir: Path):
    ms = parse_docx(sample_docx, sample_figures_dir)
    assert "Treatment X" in ms.title


def test_parse_docx_long_heading_styled_paragraph_kept_as_content(tmp_path: Path):
    """A full prose paragraph mis-styled as Heading must stay in raw_text.

    Seen in real manuscripts: entire abstracts styled as headings, which made
    them invisible to every checker.
    """
    from docx import Document

    doc = Document()
    doc.add_heading("Abstract", level=2)
    long_text = (
        "The transition from laparoscopic to robotic surgery raises safety "
        "concerns during the learning curve, particularly when complex cases "
        "are preferentially selected for the robotic platform across centers."
    )
    doc.add_heading(long_text, level=3)  # mis-styled content
    path = tmp_path / "ms.docx"
    doc.save(str(path))

    ms = parse_docx(path, None)
    assert long_text in ms.raw_text
    assert all(len(s.heading.split()) <= 25 for s in ms.sections)


def test_parse_docx_title_skips_metadata_lines(tmp_path: Path):
    from docx import Document

    doc = Document()
    doc.add_paragraph("Article Type: Observational Study")
    doc.add_paragraph("Running head: NMF Decomposition")
    doc.add_heading("Quantifying the Relative Contributions of Case Complexity", level=1)
    doc.add_paragraph("Body text.")
    path = tmp_path / "ms.docx"
    doc.save(str(path))

    ms = parse_docx(path, None)
    assert ms.title == "Quantifying the Relative Contributions of Case Complexity"


def test_parse_docx_references_end_at_next_heading(tmp_path: Path):
    """Sections after References (e.g. Figure Legends) must not leak into it."""
    from docx import Document

    doc = Document()
    doc.add_heading("Introduction", level=2)
    doc.add_paragraph("Intro text [1].")
    doc.add_heading("References", level=2)
    doc.add_paragraph("Smith J. (2020). A study. Journal, 1(1), 1-10.")
    doc.add_heading("Figure Legends", level=2)
    doc.add_paragraph("Figure 1. Mechanism of action.")
    path = tmp_path / "ms.docx"
    doc.save(str(path))

    ms = parse_docx(path, None)
    assert ms.reference_section is not None
    assert "Smith" in ms.reference_section
    assert "Mechanism of action" not in ms.reference_section


def test_parse_docx_body_text_excludes_reference_list(tmp_path: Path):
    """Citation scans use body_text — journal volume(issue) patterns in the
    reference list (e.g. '2022;101(27)') must not appear as phantom citations."""
    from docx import Document

    doc = Document()
    doc.add_heading("Introduction", level=2)
    doc.add_paragraph("A claim [1].")
    doc.add_heading("References", level=2)
    doc.add_paragraph("Smith J. A study. Medicine. 2022;101(27):e29325.")
    path = tmp_path / "ms.docx"
    doc.save(str(path))

    ms = parse_docx(path, None)
    assert "A claim [1]." in ms.body_text
    assert "101(27)" not in ms.body_text

    from sub_checker.tools.manuscript_tools import extract_citation_numbers

    assert extract_citation_numbers(ms.body_text) == {1}


def test_parse_docx_table_content_visible(tmp_path: Path):
    """Table cell text must reach raw_text and the enclosing section."""
    import docx

    doc = docx.Document()
    doc.add_heading("Results", level=1)
    doc.add_paragraph("As shown in Table 1, outcomes differed.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Group"
    table.cell(0, 1).text = "Survival rate"
    table.cell(1, 0).text = "Treatment X"
    table.cell(1, 1).text = "85.3%"
    doc.add_heading("Discussion", level=1)
    doc.add_paragraph("We discuss the results here.")
    path = tmp_path / "with_table.docx"
    doc.save(str(path))

    ms = parse_docx(path, None)
    assert "Survival rate" in ms.raw_text
    assert "85.3%" in ms.raw_text
    # Table rows are attached to the section they appear in
    results = next(s for s in ms.sections if s.heading == "Results")
    assert any("85.3%" in p.text for p in results.paragraphs)
    # And not leaked into the next section
    discussion = next(s for s in ms.sections if s.heading == "Discussion")
    assert not any("85.3%" in p.text for p in discussion.paragraphs)


def test_table_repeated_cell_values_preserved(tmp_path: Path):
    """Coincidentally-equal adjacent cells (non-merged) must NOT be collapsed."""
    import docx

    doc = docx.Document()
    doc.add_heading("Results", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.cell(0, 0).text = "Arm A"
    table.cell(0, 1).text = "10"
    table.cell(0, 2).text = "10"  # baseline N and follow-up N both 10 — a real data cell
    path = tmp_path / "repeated.docx"
    doc.save(str(path))

    ms = parse_docx(path, None)
    assert "Arm A | 10 | 10" in ms.raw_text


def test_table_merged_cells_collapsed(tmp_path: Path):
    """A genuine horizontal merge still collapses to a single cell."""
    import docx

    doc = docx.Document()
    doc.add_heading("Results", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.cell(0, 0).text = "Header"
    merged = table.cell(0, 1).merge(table.cell(0, 2))
    merged.text = "Spanning"
    path = tmp_path / "merged.docx"
    doc.save(str(path))

    ms = parse_docx(path, None)
    assert "Header | Spanning" in ms.raw_text
    assert "Spanning | Spanning" not in ms.raw_text
