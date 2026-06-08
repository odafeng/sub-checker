"""Shared test fixtures."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from sub_checker.config import Config
from sub_checker.models import Manuscript


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    """Create a sample .docx file with known content for testing."""
    doc = Document()

    # Title
    doc.add_heading("Effects of Treatment X on Disease Y: A Randomized Trial", level=1)

    # Abstract
    doc.add_heading("Abstract", level=2)
    doc.add_paragraph(
        "Background: Disease Y affects millions worldwide. "
        "Treatment X has shown promissing results in preclinical studies. "  # intentional typo
        "Methods: We conducted a randomized controlled trial with 200 patients. "
        "Results: Treatment X reduced symptoms by 45% (p<0.001). "
        "Conclusions: Treatment X is effective for Disease Y."
    )

    # Keywords
    doc.add_heading("Keywords", level=2)
    doc.add_paragraph("Disease Y, Treatment X, randomized controlled trial, efficacy")

    # Introduction
    doc.add_heading("Introduction", level=2)
    doc.add_paragraph(
        "Disease Y is a major public health concern affecting approximately "
        "10 million people globally (Smith, 2020). Current treatments are limited "
        "and often ineffective (Jones et al., 2019)."
    )
    doc.add_paragraph(
        "Treatment X, a novel therapeutic approach, has demonstrated efficacy "
        "in animal models (Brown & Lee, 2021). Figure 1 shows the mechanism of action. "
        "As shown in Table 1, previous studies have reported varying results."
    )

    # Methods
    doc.add_heading("Methods", level=2)
    doc.add_paragraph(
        "We enrolled 200 patients with confirmed Disease Y from three hospitals. "
        "Patients were randomized 1:1 to receive Treatment X or placebo. "
        "The primary endpoint was symptom reduction at 12 weeks (Figure 2)."
    )

    # Results
    doc.add_heading("Results", level=2)
    doc.add_paragraph(
        "Treatment X reduced symptoms by 45% compared to 10% in the placebo group "
        "(p<0.001, Figure 3). As shown in Table 2, adverse events were minimal. "
        "These findings are consistent with previous reports (Wilson, 2022)."
    )

    # Discussion
    doc.add_heading("Discussion", level=2)
    doc.add_paragraph(
        "Our results demonstrate that Treatment X is highly effective for Disease Y. "
        "The 45% reduction in symptoms exceeds the 30% threshold considered clinically "
        "meaningful (Garcia, 2018). However, our study has limitations including "
        "the relatively short follow-up period."
    )

    # References
    doc.add_heading("References", level=2)
    doc.add_paragraph(
        "Smith J. (2020). Epidemiology of Disease Y. Journal of Medicine, 15(3), 100-110."
    )
    doc.add_paragraph(
        "Jones A, et al. (2019). Current treatments for Disease Y. Lancet, 394, 50-60."
    )
    doc.add_paragraph(
        "Brown K, Lee M. (2021). Treatment X in animal models. Nature Medicine, 27, 200-208."
    )
    doc.add_paragraph("Garcia R. (2018). Clinical significance thresholds. BMJ, 360, k1234.")
    doc.add_paragraph("Wilson P. (2022). Treatment X pilot study. JAMA, 328(5), 450-455.")

    docx_path = tmp_path / "manuscript.docx"
    doc.save(str(docx_path))
    return docx_path


@pytest.fixture
def sample_figures_dir(tmp_path: Path) -> Path:
    """Create a figures directory with some sample files."""
    fig_dir = tmp_path / "figures"
    fig_dir.mkdir()
    (fig_dir / "Figure1.png").write_bytes(b"fake png data")
    (fig_dir / "Figure2.png").write_bytes(b"fake png data")
    # Note: Figure3.png is NOT created (referenced in text but missing)
    return fig_dir


@pytest.fixture
def sample_manuscript(sample_docx: Path, sample_figures_dir: Path) -> Manuscript:
    """Parse sample docx into Manuscript model."""
    from sub_checker.parsers.docx_parser import parse_docx

    return parse_docx(sample_docx, sample_figures_dir)


@pytest.fixture
def test_config() -> Config:
    """Config with COT logging disabled for test isolation."""
    return Config(cot_dir="disabled")
