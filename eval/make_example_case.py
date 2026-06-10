"""Generate a synthetic example case so the eval runner can be tried
without any private manuscript.

Usage: python eval/make_example_case.py
Creates eval/cases/example/ (gitignored) with a manuscript containing
known seeded issues and a matching golden.yaml.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

CASE_DIR = Path(__file__).parent / "cases" / "example"

GOLDEN_YAML = """\
journal: null
checkers: [typo_grammar, figure_table, citation_exist]
lang: en

expected:
  - id: typo-promissing
    description: "intentional misspelling in the abstract"
    checker: typo_grammar
    keywords: ["promissing"]
  - id: missing-figure-3
    description: "Figure 3 is referenced but the file does not exist"
    checker: figure_table
    keywords: ["Figure 3"]

forbidden:
  - id: fp-ref-2-uncited
    description: "reference [2] IS cited; flagging it as uncited is a false positive"
    claim_type: uncited_reference
    ref_number: 2
"""


def main() -> None:
    CASE_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading("Effects of Treatment X on Disease Y: A Randomized Trial", level=1)
    doc.add_heading("Abstract", level=2)
    doc.add_paragraph(
        "Background: Disease Y affects millions worldwide. "
        "Treatment X has shown promissing results in preclinical studies. "  # seeded typo
        "Methods: We conducted a randomized controlled trial with 200 patients. "
        "Results: Treatment X reduced symptoms by 45% (p<0.001). "
        "Conclusions: Treatment X is effective for Disease Y."
    )
    doc.add_heading("Introduction", level=2)
    doc.add_paragraph(
        "Disease Y is a major public health concern [1]. "
        "Current treatments are limited and often ineffective [2]. "
        "Figure 1 shows the mechanism of action."
    )
    doc.add_heading("Methods", level=2)
    doc.add_paragraph(
        "We enrolled 200 patients with confirmed Disease Y. "
        "The primary endpoint was symptom reduction at 12 weeks (Figure 2)."
    )
    doc.add_heading("Results", level=2)
    doc.add_paragraph(
        "Treatment X reduced symptoms by 45% compared to 10% in the placebo "
        "group (p<0.001, Figure 3). Adverse events were minimal [3]."
    )
    doc.add_heading("References", level=2)
    doc.add_paragraph(
        "Smith J. (2020). Epidemiology of Disease Y. Journal of Medicine, 15(3), 100-110."
    )
    doc.add_paragraph(
        "Jones A, et al. (2019). Current treatments for Disease Y. Lancet, 394, 50-60."
    )
    doc.add_paragraph("Wilson P. (2022). Treatment X pilot study. JAMA, 328(5), 450-455.")
    doc.save(str(CASE_DIR / "manuscript.docx"))

    fig_dir = CASE_DIR / "figures"
    fig_dir.mkdir(exist_ok=True)
    (fig_dir / "Figure1.png").write_bytes(b"fake png data")
    (fig_dir / "Figure2.png").write_bytes(b"fake png data")
    # Figure3.png intentionally missing

    (CASE_DIR / "golden.yaml").write_text(GOLDEN_YAML)
    print(f"Example case created at {CASE_DIR}")
    print("Run it with: sub-check-eval --case example")


if __name__ == "__main__":
    main()
