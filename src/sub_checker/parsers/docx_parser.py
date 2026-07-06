"""Parse .docx files into Manuscript model."""

from __future__ import annotations

import re
from collections.abc import Iterator
from pathlib import Path

import docx
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph as DocxParagraph

from sub_checker.models import Manuscript, Paragraph, Section

_REFERENCE_HEADINGS = {"references", "bibliography", "works cited", "literature cited"}
_ABSTRACT_HEADINGS = {"abstract", "summary"}
# Common section headings that may appear as plain text (Normal style) in .docx
_SECTION_HEADINGS = {
    "introduction",
    "methods",
    "materials and methods",
    "results",
    "discussion",
    "conclusions",
    "conclusion",
    "acknowledgments",
    "acknowledgements",
    "disclosures",
    "funding",
    "figure legends",
    "table legends",
    "supplementary materials",
    "supplementary material",
    "appendix",
}

# Real section headings are short. Authors sometimes apply a Heading style to
# whole paragraphs (seen in real manuscripts: entire abstracts/discussion
# paragraphs styled as headings) — treating those as headings would make the
# text invisible to every checker, so anything longer is kept as content.
_MAX_HEADING_WORDS = 25

# Header lines that are submission metadata, not the manuscript title
_METADATA_LINE = re.compile(
    r"^(article\s+type|running\s+(head|title)|short\s+title|title\s+page|"
    r"word\s+count|corresponding\s+author|authors?|affiliations?|keywords?)\b",
    re.IGNORECASE,
)


def _pick_title(header_lines: list[str], sections: list[Section]) -> str:
    """Choose the manuscript title: first non-metadata header line, else first heading."""
    for line in header_lines:
        if not _METADATA_LINE.match(line):
            return line
    for section in sections:
        if not _METADATA_LINE.match(section.heading):
            return section.heading
    return header_lines[0] if header_lines else "Untitled"


def _iter_block_items(doc) -> Iterator[DocxParagraph | Table]:
    """Yield paragraphs AND tables in true document order.

    `doc.paragraphs` skips tables entirely, making table content (e.g.
    Table 1 data, headers, footnotes) invisible to every checker.
    """
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield DocxParagraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield Table(child, doc)


def _table_row_texts(table: Table) -> list[str]:
    """Flatten a table into one text line per row ('cell | cell | cell')."""
    rows = []
    for row in table.rows:
        cells: list[str] = []
        prev_tc = None
        for cell in row.cells:
            # A horizontally-merged cell exposes the SAME underlying <w:tc>
            # element at each grid position, so collapse by element identity.
            # Comparing cell TEXT instead would also drop legitimately repeated
            # values in non-merged cells (e.g. a "Arm A | 10 | 10" data row).
            if cell._tc is prev_tc:
                continue
            prev_tc = cell._tc
            ct = " ".join(cell.text.split())
            if ct:
                cells.append(ct)
        if cells:
            rows.append(" | ".join(cells))
    return rows


def parse_docx(docx_path: Path, figure_dir: Path | None = None) -> Manuscript:
    """Parse a .docx file into a Manuscript model."""
    doc = docx.Document(str(docx_path))

    paragraphs: list[Paragraph] = []
    sections: list[Section] = []
    current_section: Section | None = None
    reference_section: str | None = None
    in_references = False
    ref_lines: list[str] = []
    body_lines: list[str] = []  # Everything except the reference list
    header_lines: list[str] = []  # Text before first heading
    first_heading_seen = False

    def add_content_paragraph(text: str, is_table: bool = False) -> None:
        p = Paragraph(
            text=text,
            index=len(paragraphs),
            section=current_section.heading if current_section else None,
        )
        paragraphs.append(p)
        if current_section:
            current_section.paragraphs.append(p)
        # Table rows are never references — even when a table is laid out after
        # the "References" heading, it must not pollute the reference list
        # (which would inflate the count and fabricate "uncited" numbers).
        if in_references and not is_table:
            ref_lines.append(text)
        else:
            body_lines.append(text)

    for block in _iter_block_items(doc):
        if isinstance(block, Table):
            # Table content stays in document order, attached to the current
            # section, so checkers can see in-table inconsistencies.
            for row_text in _table_row_texts(block):
                add_content_paragraph(row_text, is_table=True)
            continue

        para = block
        text = para.text.strip()
        if not text:
            continue

        style = para.style
        style_name = ((style.name or "") if style else "").lower()
        # A heading-styled paragraph that is actually a full paragraph of prose
        # must be treated as content, or it disappears from raw_text entirely.
        is_heading = "heading" in style_name and len(text.split()) <= _MAX_HEADING_WORDS
        is_ref_heading = text.lower() in _REFERENCE_HEADINGS
        is_abstract_heading = text.lower() in _ABSTRACT_HEADINGS
        is_section_heading = text.lower() in _SECTION_HEADINGS

        if is_heading or is_ref_heading or is_abstract_heading or is_section_heading:
            first_heading_seen = True
            level = 1
            for ch in style_name:
                if ch.isdigit():
                    level = int(ch)
                    break

            # Track whether we're inside the reference list: a non-reference
            # heading (e.g. "Figure Legends" after "References") ends it.
            in_references = is_ref_heading

            current_section = Section(heading=text, level=level)
            sections.append(current_section)
            continue

        # Collect text before first heading as header
        if not first_heading_seen:
            header_lines.append(text)

        add_content_paragraph(text)

    if ref_lines:
        reference_section = "\n".join(ref_lines)

    raw_text = "\n".join(p.text for p in paragraphs)
    header_text = "\n".join(header_lines)
    body_text = "\n".join(body_lines)

    # Title: first non-metadata line before any heading; fall back to first heading
    title = _pick_title(header_lines, sections)

    return Manuscript(
        title=title,
        sections=sections,
        paragraphs=paragraphs,
        raw_text=raw_text,
        reference_section=reference_section,
        figure_dir=figure_dir,
        header_text=header_text,
        body_text=body_text,
    )
